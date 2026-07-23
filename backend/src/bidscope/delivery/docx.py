"""DOCX rendering and attachment for persisted typed reports.

Rendering is pure. The delivery service writes a deterministic object and only
then attaches its key to an existing online report row. Online report creation
belongs to :mod:`bidscope.delivery.reports`, so DOCX failure never rolls back
already-persisted report evidence.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass
from datetime import datetime
from typing import TYPE_CHECKING

from bidscope.delivery.objects import ObjectStore
from bidscope.domain.reports import Report
from bidscope.domain.types import BidScopeErrorCode
from bidscope.persistence.models import Report as ReportModel
from docx import Document
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker

if TYPE_CHECKING:
    from bidscope.delivery.reports import PersistedReport


#: Bumped when the rendered output format changes. It participates in the
#: deterministic DOCX object key while the report row keeps its online key.
RENDERER_VERSION = "docx-v1"


class DeliveryError(Exception):
    """A delivery/export failure carrying a bounded error code."""

    def __init__(
        self,
        message: str,
        code: BidScopeErrorCode = BidScopeErrorCode.DELIVERY_ERROR,
        cause: BaseException | None = None,
    ) -> None:
        super().__init__(message)
        self.code = code
        self.cause = cause


@dataclass(frozen=True)
class ExportRecord:
    """One deterministic DOCX attachment for a persisted online report."""

    export_key: str
    object_key: str
    report_id: str
    generated_at: datetime


def _export_key(report_id: str) -> str:
    """Derive DOCX idempotency from persisted report identity and renderer."""
    return f"{RENDERER_VERSION}:{report_id}"


def _sanitize_filename(name: str) -> str:
    """Strip characters unsafe inside an object key or DOCX filename."""
    sanitized = re.sub(r"[^A-Za-z0-9._-]", "_", name)
    sanitized = sanitized.strip(".")
    return sanitized or "report"


def _object_key(report_id: str) -> str:
    """Build a deterministic object key from report identity and renderer."""
    safe = _sanitize_filename(report_id)
    return f"reports/{safe}/{RENDERER_VERSION}/bidscope-{safe}.docx"


def render_report(report: Report) -> bytes:
    """Render a typed report to DOCX bytes without any external effects."""
    document = Document()
    document.add_heading("BidScope Report", level=0)

    document.add_paragraph(f"Generated at: {report.generated_at.isoformat()}")
    if report.freshness_window:
        document.add_paragraph(f"Freshness window: {report.freshness_window}")

    document.add_heading("Query Conditions", level=1)
    conditions_table = document.add_table(rows=1, cols=2)
    hdr = conditions_table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"
    for key, value in report.query_conditions.items():
        row = conditions_table.add_row().cells
        row[0].text = key
        row[1].text = str(value)

    if report.source_availability:
        document.add_heading("Source Availability", level=1)
        for source in report.source_availability:
            document.add_paragraph(source)

    if report.completeness_warning:
        document.add_heading("Completeness Warning", level=1)
        document.add_paragraph(report.completeness_warning)

    document.add_heading("Opportunities", level=1)
    for idx, item in enumerate(report.items, start=1):
        document.add_heading(f"{idx}. {item.title}", level=2)

        if item.known_fields:
            document.add_paragraph("Known fields:")
            kf_table = document.add_table(rows=1, cols=2)
            kf_hdr = kf_table.rows[0].cells
            kf_hdr[0].text = "Field"
            kf_hdr[1].text = "Value"
            for key, value in item.known_fields.items():
                row = kf_table.add_row().cells
                row[0].text = key
                row[1].text = str(value)

        if item.unknown_fields:
            document.add_paragraph("未知字段 (unknown fields): " + ", ".join(item.unknown_fields))

        if item.relevance_reason:
            document.add_paragraph(f"Relevance: {item.relevance_reason}")
        if item.risk_note:
            document.add_paragraph(f"Risk: {item.risk_note}")

        if item.citations:
            document.add_paragraph("Evidence:")
            for cidx, citation in enumerate(item.citations, start=1):
                label = citation.label or citation.evidence_id
                document.add_paragraph(f"  [{cidx}] {label}")

        if item.claims:
            document.add_paragraph("Claims:")
            for claim in item.claims:
                document.add_paragraph(f"  - {claim.text}")

    document.add_heading("Appendix", level=1)
    document.add_paragraph(f"Renderer version: {RENDERER_VERSION}")
    document.add_paragraph(f"Report run ID: {report.run_id}")
    document.add_paragraph(
        "This document was generated automatically and reflects the "
        "evidence available at generation time."
    )

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class ReportDelivery:
    """Attach idempotent DOCX objects to already-persisted online reports."""

    def __init__(
        self,
        store: ObjectStore,
        session_factory: async_sessionmaker[AsyncSession],
    ) -> None:
        self.store = store
        self.session_factory = session_factory

    async def export_report(self, persisted: PersistedReport) -> ExportRecord:
        """Render and attach a DOCX without creating or replacing a report row."""
        async with self.session_factory() as session:
            row = await session.get(ReportModel, persisted.id)
            if row is None:
                raise DeliveryError("online report is not persisted")
            if row.docx_object_key:
                try:
                    attached_object_exists = self.store.exists(row.docx_object_key)
                except Exception as exc:
                    raise DeliveryError(
                        f"DOCX storage failed for report {persisted.id}",
                        code=BidScopeErrorCode.DELIVERY_ERROR,
                        cause=exc,
                    ) from exc
                if attached_object_exists:
                    return ExportRecord(
                        export_key=_export_key(str(row.id)),
                        object_key=row.docx_object_key,
                        report_id=str(row.id),
                        generated_at=row.generated_at,
                    )
                object_key = row.docx_object_key
            else:
                object_key = _object_key(persisted.id)

        data = render_report(persisted.report)
        try:
            self.store.put_bytes(object_key, data)
        except Exception as exc:
            raise DeliveryError(
                f"DOCX storage failed for report {persisted.id}",
                code=BidScopeErrorCode.DELIVERY_ERROR,
                cause=exc,
            ) from exc

        async with self.session_factory() as session:
            row = await session.get(ReportModel, persisted.id)
            if row is None:
                raise DeliveryError("online report disappeared before DOCX attachment")
            if row.docx_object_key is None:
                row.docx_object_key = object_key
                await session.commit()
            return ExportRecord(
                export_key=_export_key(str(row.id)),
                object_key=row.docx_object_key,
                report_id=str(row.id),
                generated_at=row.generated_at,
            )


__all__ = [
    "DeliveryError",
    "ExportRecord",
    "RENDERER_VERSION",
    "ReportDelivery",
    "render_report",
]
