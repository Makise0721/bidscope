"""Integration coverage for online report persistence and DOCX attachment."""

from __future__ import annotations

import asyncio
import io
import uuid
from datetime import UTC, datetime

import docx
import pytest
import sqlalchemy as sa
from bidscope.api.dependencies import RunService
from bidscope.api.routes.reports import get_report
from bidscope.config import Settings
from bidscope.delivery.docx import DeliveryError
from bidscope.delivery.objects import LocalObjectStore
from bidscope.delivery.reports import ReportPersistence
from bidscope.domain.notices import NoticeEvidence as DomainEvidence
from bidscope.domain.reports import Report, ReportCitation, ReportClaim, ReportItem
from bidscope.observability import MetricsRegistry
from bidscope.persistence.models import (
    CanonicalNotice,
    NoticeEvidence,
    NoticeVersion,
    QueryRun,
    ReportClaimCitation,
    SourceNotice,
)
from bidscope.persistence.models import (
    Report as ReportModel,
)
from bidscope.persistence.models import (
    ReportCitation as ReportCitationModel,
)
from bidscope.persistence.models import (
    ReportClaim as ReportClaimModel,
)
from bidscope.persistence.models import (
    ReportItem as ReportItemModel,
)
from sqlalchemy.ext.asyncio import AsyncSession


async def _count(session: AsyncSession, model: type) -> int:
    return (await session.scalar(sa.select(sa.func.count()).select_from(model))) or 0


async def _seed_report_inputs(session_factory) -> tuple[Report, dict[str, DomainEvidence]]:
    run_id = str(uuid.uuid4())
    now = datetime(2026, 7, 18, 12, 0, tzinfo=UTC)
    async with session_factory() as session:
        canonical = CanonicalNotice(title="四川智算中心服务器采购招标公告")
        session.add(canonical)
        await session.flush()
        source = SourceNotice(
            canonical_notice_id=canonical.id,
            source="synthetic_demo",
            external_id="report-delivery-demo",
            source_url="https://example.invalid/report-delivery-demo",
            first_seen_at=now,
            latest_seen_at=now,
            content_hash="notice-hash",
            title="四川智算中心服务器采购招标公告",
        )
        session.add(source)
        await session.flush()
        version = NoticeVersion(
            source_notice_id=source.id,
            payload_object_key="snapshots/demo/report-delivery-demo",
            capture_kind="synthetic_demo",
            parser_version="demo-v1",
            content_hash="version-hash",
            title=source.title,
            region="四川",
        )
        session.add(version)
        await session.flush()
        evidence = NoticeEvidence(
            notice_version_id=version.id,
            text="预算800万元。",
            start=0,
            end=8,
            span_hash="evidence-001",
        )
        session.add(evidence)
        second_evidence = NoticeEvidence(
            notice_version_id=version.id,
            text="交付地点为四川省。",
            start=8,
            end=17,
            span_hash="evidence-002",
        )
        session.add(second_evidence)
        session.add(QueryRun(
            id=run_id,
            run_key=run_id,
            status="running",
            user_request="test",
        ))
        await session.commit()

    domain_evidence = {
        evidence.span_hash: DomainEvidence(
            notice_version_id=str(version.id),
            text=evidence.text,
            start=evidence.start,
            end=evidence.end,
            span_hash=evidence.span_hash,
        ),
        second_evidence.span_hash: DomainEvidence(
            notice_version_id=str(version.id),
            text=second_evidence.text,
            start=second_evidence.start,
            end=second_evidence.end,
            span_hash=second_evidence.span_hash,
        ),
    }
    report = Report(
        run_id=run_id,
        generated_at=now,
        query_conditions={"region": "四川", "budget": ">=500万"},
        completeness_warning="部分数据源暂不可用。",
        items=[
            ReportItem(
                notice_id=str(version.id),
                title=version.title or "",
                known_fields={"source_url": source.source_url},
                unknown_fields=["deadline"],
                relevance_reason="服务器采购",
                risk_note="核验截止时间",
                citations=[ReportCitation(evidence_id=evidence.span_hash, label="预算金额证据")],
                claims=[ReportClaim(text=evidence.text, citation_ids=[evidence.span_hash])],
            )
        ],
    )
    return report, domain_evidence


@pytest.mark.asyncio
async def test_persist_online_report_stores_items_claims_evidence_and_docx(
    session_factory,
    tmp_path,
) -> None:
    report, evidence_by_hash = await _seed_report_inputs(session_factory)
    persistence = ReportPersistence(session_factory, LocalObjectStore(tmp_path / "objects"))

    persisted = await persistence.persist_online_report(report, evidence_by_hash)
    export = await persistence.export_docx(persisted)

    assert persistence.store.exists(export.object_key)
    document = docx.Document(io.BytesIO(persistence.store.get_bytes(export.object_key)))
    assert isinstance(document, docx.document.Document)
    async with session_factory() as session:
        assert await _count(session, ReportModel) == 1
        assert await _count(session, ReportItemModel) == 1
        assert await _count(session, ReportCitationModel) == 1
        assert await _count(session, ReportClaimModel) == 1
        assert await _count(session, ReportClaimCitation) == 1
        stored = await session.get(ReportModel, persisted.id)
        assert stored is not None
        assert stored.docx_object_key == export.object_key


@pytest.mark.asyncio
async def test_report_delivery_duration_metrics_record_success_for_online_and_docx(
    session_factory,
    tmp_path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Both delivery stages must use the registry's bounded outcome label."""
    registry = MetricsRegistry()
    monkeypatch.setattr("bidscope.delivery.reports.METRICS_REGISTRY", registry)
    report, evidence_by_hash = await _seed_report_inputs(session_factory)
    persistence = ReportPersistence(session_factory, LocalObjectStore(tmp_path / "objects"))

    persisted = await persistence.persist_online_report(report, evidence_by_hash)
    await persistence.export_docx(persisted)

    assert (
        'bidscope_report_delivery_duration_seconds_count{outcome="success"} 2'
        in registry.render_prometheus()
    )


@pytest.mark.asyncio
async def test_replay_keeps_one_online_report_and_one_docx_attachment(
    session_factory,
    tmp_path,
) -> None:
    report, evidence_by_hash = await _seed_report_inputs(session_factory)
    persistence = ReportPersistence(session_factory, LocalObjectStore(tmp_path / "objects"))

    first = await persistence.persist_online_report(report, evidence_by_hash)
    second = await persistence.persist_online_report(report, evidence_by_hash)
    first_export = await persistence.export_docx(first)
    second_export = await persistence.export_docx(second)

    assert first.id == second.id
    assert first_export.object_key == second_export.object_key
    async with session_factory() as session:
        assert await _count(session, ReportModel) == 1
        assert await _count(session, ReportItemModel) == 1


@pytest.mark.asyncio
async def test_mismatched_evidence_aborts_online_report_transaction(
    session_factory, tmp_path
) -> None:
    report, evidence_by_hash = await _seed_report_inputs(session_factory)
    evidence = next(iter(evidence_by_hash.values()))
    foreign = evidence.model_copy(update={"notice_version_id": str(uuid.uuid4())})
    persistence = ReportPersistence(session_factory, LocalObjectStore(tmp_path / "objects"))

    with pytest.raises(DeliveryError):
        await persistence.persist_online_report(report, {foreign.span_hash: foreign})

    async with session_factory() as session:
        assert await _count(session, ReportModel) == 0


@pytest.mark.asyncio
async def test_docx_failure_keeps_online_report_and_later_retry_works(
    session_factory, tmp_path
) -> None:
    class FailingObjectStore:
        def put_bytes(self, key: str, data: bytes) -> str:
            raise OSError("object store unavailable")

        def get_bytes(self, key: str) -> bytes:
            raise FileNotFoundError(key)

        def exists(self, key: str) -> bool:
            return False

    report, evidence_by_hash = await _seed_report_inputs(session_factory)
    report = report.model_copy(update={"source_availability": ["ccgp", "sichuan"]})
    failed = ReportPersistence(session_factory, FailingObjectStore())
    persisted = await failed.persist_online_report(report, evidence_by_hash)

    with pytest.raises(DeliveryError):
        await failed.export_docx(persisted)

    async with session_factory() as session:
        stored = await session.get(ReportModel, persisted.id)
        assert stored is not None
        assert stored.docx_object_key is None
        assert stored.source_availability == ["ccgp", "sichuan"]

    retry = ReportPersistence(session_factory, LocalObjectStore(tmp_path / "objects"))
    replayed = await retry.persist_online_report(report, evidence_by_hash)
    assert replayed.report.source_availability == ["ccgp", "sichuan"]

    export = await retry.export_docx(replayed)
    assert retry.store.exists(export.object_key)
    document = docx.Document(io.BytesIO(retry.store.get_bytes(export.object_key)))
    assert "ccgp" in [paragraph.text for paragraph in document.paragraphs]
    assert "sichuan" in [paragraph.text for paragraph in document.paragraphs]

    service = RunService(
        session_factory,
        graph=object(),
        object_store=retry.store,
        settings=Settings(app_mode="test"),
    )
    assert (await get_report(report.run_id, service))["source_availability"] == ["ccgp", "sichuan"]
    async with session_factory() as session:
        assert await _count(session, ReportModel) == 1


@pytest.mark.asyncio
async def test_report_replay_and_api_preserve_claim_and_citation_order(
    session_factory,
    tmp_path,
) -> None:
    report, evidence_by_hash = await _seed_report_inputs(session_factory)
    item = report.items[0].model_copy(update={
        "citations": [
            ReportCitation(evidence_id="evidence-002", label="delivery location"),
            ReportCitation(evidence_id="evidence-001", label="budget"),
        ],
        "claims": [
            ReportClaim(
                text="预算和交付地点已确认。",
                citation_ids=["evidence-002", "evidence-001"],
            ),
            ReportClaim(text="预算为800万元。", citation_ids=["evidence-001"]),
        ],
    })
    report = report.model_copy(update={"items": [item]})
    persistence = ReportPersistence(session_factory, LocalObjectStore(tmp_path / "objects"))

    await persistence.persist_online_report(report, evidence_by_hash)
    replayed = await persistence.persist_online_report(report, evidence_by_hash)

    replayed_item = replayed.report.items[0]
    assert [citation.evidence_id for citation in replayed_item.citations] == [
        "evidence-002",
        "evidence-001",
    ]
    assert [claim.text for claim in replayed_item.claims] == [
        "预算和交付地点已确认。",
        "预算为800万元。",
    ]
    assert replayed_item.claims[0].citation_ids == ["evidence-002", "evidence-001"]

    service = RunService(
        session_factory,
        graph=object(),
        object_store=persistence.store,
        settings=Settings(app_mode="test"),
    )
    body = await get_report(report.run_id, service)
    api_item = body["items"][0]
    assert [citation["span_hash"] for citation in api_item["citations"]] == [
        "evidence-002",
        "evidence-001",
    ]
    assert [claim["text"] for claim in api_item["claims"]] == [
        "预算和交付地点已确认。",
        "预算为800万元。",
    ]
    assert api_item["claims"][0]["citation_ids"] == ["evidence-002", "evidence-001"]


@pytest.mark.asyncio
async def test_persistence_deduplicates_bypassed_claim_citation_ids_in_first_seen_order(
    session_factory,
    tmp_path,
) -> None:
    report, evidence_by_hash = await _seed_report_inputs(session_factory)
    item = report.items[0].model_copy(update={
        "claims": [
            ReportClaim.model_construct(
                text="预算和交付地点已确认。",
                citation_ids=["evidence-002", "evidence-001", "evidence-002"],
            )
        ],
    })
    report = report.model_copy(update={"items": [item]})
    persistence = ReportPersistence(session_factory, LocalObjectStore(tmp_path / "objects"))

    persisted = await persistence.persist_online_report(report, evidence_by_hash)
    replayed = await persistence.load_online_report(report.run_id)

    assert replayed is not None
    assert replayed.id == persisted.id
    assert replayed.report.items[0].claims[0].citation_ids == ["evidence-002", "evidence-001"]
    async with session_factory() as session:
        assert await _count(session, ReportClaimCitation) == 2


@pytest.mark.asyncio
async def test_database_enforces_one_report_per_non_null_run_id(session_factory) -> None:
    run_id = str(uuid.uuid4())
    generated_at = datetime(2026, 7, 18, 12, 0, tzinfo=UTC)
    async with session_factory() as session:
        session.add(QueryRun(
            id=run_id,
            run_key=run_id,
            status="running",
            user_request="test",
        ))
        session.add(ReportModel(
            run_id=run_id,
            export_key="external:one",
            generated_at=generated_at,
        ))
        await session.commit()

        session.add(ReportModel(
            run_id=run_id,
            export_key="external:two",
            generated_at=generated_at,
        ))
        with pytest.raises(sa.exc.IntegrityError):
            await session.commit()
        await session.rollback()

        session.add_all([
            ReportModel(export_key="external:legacy-one", generated_at=generated_at),
            ReportModel(export_key="external:legacy-two", generated_at=generated_at),
        ])
        await session.commit()


@pytest.mark.asyncio
async def test_concurrent_persistence_creates_one_report_for_a_run(
    session_factory,
    tmp_path,
) -> None:
    report, evidence_by_hash = await _seed_report_inputs(session_factory)
    persistence = ReportPersistence(session_factory, LocalObjectStore(tmp_path / "objects"))
    original_find_by_run_id = persistence._find_by_run_id
    first_lookup_complete = asyncio.Event()
    both_lookups_complete = asyncio.Event()
    lookups = 0

    async def synchronized_find_by_run_id(session, run_id):
        nonlocal lookups
        if lookups < 2:
            lookups += 1
            first_lookup_complete.set()
            if lookups == 2:
                both_lookups_complete.set()
            await both_lookups_complete.wait()
        return await original_find_by_run_id(session, run_id)

    persistence._find_by_run_id = synchronized_find_by_run_id  # type: ignore[method-assign]
    first_task = asyncio.create_task(persistence.persist_online_report(report, evidence_by_hash))
    await first_lookup_complete.wait()
    second_task = asyncio.create_task(persistence.persist_online_report(report, evidence_by_hash))
    first, second = await asyncio.gather(first_task, second_task)

    assert first.id == second.id
    async with session_factory() as session:
        assert await _count(session, ReportModel) == 1
