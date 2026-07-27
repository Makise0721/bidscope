"""Integration tests for the runs API surface.

Covers ``POST /api/runs``, ``GET /api/runs/{id}``, ``POST /api/runs/{id}/confirm``,
``POST /api/runs/{id}/retry`` and the report endpoints.

The confirm/retry state-machine is the critical contract:

* ``confirm`` succeeds (200) only when the run is ``awaiting_confirmation``;
  otherwise it returns HTTP 409.
* ``retry`` succeeds (200) only when the run is ``retryable``; otherwise 409.

Only synthetic-demo data flows through the demo graph; no network access occurs.
"""

from __future__ import annotations

import asyncio
import io
import time
from datetime import UTC, datetime
from pathlib import Path

import docx
import pytest
import sqlalchemy as sa
from bidscope.config import Settings, get_settings
from bidscope.delivery.docx import DeliveryError
from bidscope.delivery.objects import LocalObjectStore
from bidscope.delivery.reports import ReportPersistence
from bidscope.domain.reports import Report
from bidscope.main import create_app
from bidscope.persistence.models import QueryRun
from bidscope.persistence.models import Report as ReportModel
from fastapi.testclient import TestClient

SCHEDULED_QUERY = (
    "每周一上午 9 点，汇总近 7 天四川和重庆与「智算中心、服务器」有关、"
    "预算 500 万以上的招标信息。"
)
NON_SCHEDULED_QUERY = "查询四川省最近的服务器招标信息。"
PRODUCTION_ADMIN_TOKEN = "test-admin-token-012345678901234567890123"


def _poll_status(client: TestClient, run_id: str, expected: str, timeout: float = 15.0) -> dict:
    """Poll GET /api/runs/{id} until status matches ``expected`` or timeout."""
    deadline = time.monotonic() + timeout
    last = None
    while time.monotonic() < deadline:
        response = client.get(f"/api/runs/{run_id}")
        assert response.status_code == 200, response.text
        last = response.json()
        if last["status"] == expected:
            return last
        time.sleep(0.05)
    raise AssertionError(f"run {run_id} never reached status {expected!r}; last={last}")


def test_create_run_returns_pending_then_completes(demo_client: TestClient) -> None:
    """POST /api/runs stores pending, schedules the executor, and the run progresses."""
    client = demo_client

    response = client.post("/api/runs", json={"user_request": NON_SCHEDULED_QUERY})
    assert response.status_code == 201, response.text
    body = response.json()
    assert "id" in body
    assert body["status"] == "pending"
    run_id = body["id"]

    # The background executor drives the run to a terminal/confirmation state.
    final = _poll_status(client, run_id, "completed")
    assert final["id"] == run_id


def test_create_run_rejects_empty_request(demo_client: TestClient) -> None:
    """An empty user request is a client error, not a server error."""
    response = demo_client.post("/api/runs", json={"user_request": "   "})
    assert response.status_code == 422


def test_get_run_returns_run(demo_client: TestClient) -> None:
    """GET /api/runs/{id} returns the stored run."""
    client = demo_client

    created = client.post("/api/runs", json={"user_request": NON_SCHEDULED_QUERY}).json()
    response = client.get(f"/api/runs/{created['id']}")
    assert response.status_code == 200
    assert response.json()["id"] == created["id"]


def test_get_run_not_found(demo_client: TestClient) -> None:
    """An unknown id returns 404."""
    response = demo_client.get("/api/runs/00000000-0000-0000-0000-000000000000")
    assert response.status_code == 404


def test_confirm_succeeds_when_awaiting_confirmation(demo_client: TestClient) -> None:
    """A scheduled query pauses for confirmation; approving resumes and completes it."""
    client = demo_client

    created = client.post("/api/runs", json={"user_request": SCHEDULED_QUERY}).json()
    run_id = created["id"]

    # The scheduled query pauses at confirmation; the executor updates the status.
    _poll_status(client, run_id, "awaiting_confirmation")

    response = client.post(f"/api/runs/{run_id}/confirm", json={"action": "approve"})
    assert response.status_code == 200, response.text

    # Resuming the graph runs it through retrieval to completion.
    final = _poll_status(client, run_id, "completed")
    assert final["status"] == "completed"


def test_confirm_returns_409_unless_awaiting_confirmation(demo_client: TestClient) -> None:
    """confirm rejects any run that is not awaiting confirmation."""
    client = demo_client

    # A completed run is not awaiting confirmation.
    created = client.post("/api/runs", json={"user_request": NON_SCHEDULED_QUERY}).json()
    run_id = created["id"]
    _poll_status(client, run_id, "completed")

    response = client.post(f"/api/runs/{run_id}/confirm", json={"action": "approve"})
    assert response.status_code == 409


def test_retry_returns_409_unless_retryable(demo_client: TestClient) -> None:
    """retry rejects any run that is not retryable."""
    client = demo_client

    created = client.post("/api/runs", json={"user_request": NON_SCHEDULED_QUERY}).json()
    run_id = created["id"]
    _poll_status(client, run_id, "completed")

    response = client.post(f"/api/runs/{run_id}/retry")
    assert response.status_code == 409


def test_docx_retry_exports_persisted_report_without_running_graph(tmp_path: Path) -> None:
    """DOCX retry only exports the existing online report and is idempotent."""
    class FailingObjectStore:
        def put_bytes(self, key: str, data: bytes) -> str:
            raise OSError("object store unavailable")

        def get_bytes(self, key: str) -> bytes:
            raise FileNotFoundError(key)

        def exists(self, key: str) -> bool:
            return False

    class CountingGraph:
        def __init__(self) -> None:
            self.graph_calls = 0
            self.retrieval_calls = 0

    settings = Settings(
        app_mode="production",
        database_url=get_settings().database_url,
        checkpoint_database_url=get_settings().checkpoint_database_url,
        real_model_enabled=False,
        admin_token=PRODUCTION_ADMIN_TOKEN,
        object_store_type="s3",
        s3_endpoint="http://minio:9000",
        s3_region="us-east-1",
        s3_bucket="bidscope-test",
        s3_access_key="test-access",
        s3_secret_key="test-secret",
        allowed_origins=["https://bidscope.test"],
        trusted_hosts=["bidscope.test"],
        external_scheme="https",
        object_store_root=str(tmp_path / "objects"),
    )
    headers = {"X-Admin-Token": PRODUCTION_ADMIN_TOKEN}
    with TestClient(create_app(settings=settings)) as client:
        service = client.app.state.run_service
        service.object_store = LocalObjectStore(tmp_path / "objects")
        graph = CountingGraph()
        service.graph = graph

        async def _persist_failed_report() -> tuple[str, str]:
            run_id, created = await service.create_run("DOCX retry test")
            assert created is True
            report = Report(
                run_id=run_id,
                generated_at=datetime(2026, 7, 18, 12, 0, tzinfo=UTC),
                query_conditions={"region": "四川"},
                source_availability=["ccgp", "sichuan"],
            )
            failed = ReportPersistence(service.session_factory, FailingObjectStore())
            persisted = await failed.persist_online_report(report, {})
            with pytest.raises(DeliveryError):
                await failed.export_docx(persisted)
            async with service.session_factory() as session:
                stored = await session.get(ReportModel, persisted.id)
                assert stored is not None
                assert stored.docx_object_key is None
            return run_id, persisted.id

        assert client.portal is not None
        run_id, report_id = client.portal.call(_persist_failed_report)
        service.object_store = LocalObjectStore(tmp_path / "recovered-objects")

        first = client.post(f"/api/reports/{run_id}/docx/retry", headers=headers)
        second = client.post(f"/api/reports/{run_id}/docx/retry", headers=headers)

        assert first.status_code == 200, first.text
        assert second.status_code == 200, second.text
        assert first.json()["report_id"] == report_id
        assert first.json()["docx_object_key"] == second.json()["docx_object_key"]
        document = docx.Document(io.BytesIO(service.object_store.get_bytes(
            first.json()["docx_object_key"]
        )))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        assert "ccgp" in paragraphs
        assert "sichuan" in paragraphs
        assert graph.graph_calls == 0
        assert graph.retrieval_calls == 0

        async def _report_rows() -> list[ReportModel]:
            async with service.session_factory() as session:
                result = await session.execute(
                    sa.select(ReportModel).where(ReportModel.run_id == run_id)
                )
                return list(result.scalars())

        rows = client.portal.call(_report_rows)
        assert len(rows) == 1
        assert rows[0].docx_object_key == first.json()["docx_object_key"]
        assert client.post(
            "/api/reports/00000000-0000-0000-0000-000000000000/docx/retry",
            headers=headers,
        ).status_code == 404


def test_docx_retry_recreates_a_missing_attached_object_with_the_same_key(tmp_path: Path) -> None:
    """A retry repairs a missing object without creating a new report or key."""
    class CountingObjectStore(LocalObjectStore):
        def __init__(self, root: Path) -> None:
            super().__init__(root)
            self.put_calls = 0
            self.fail_writes = False

        def put_bytes(self, key: str, data: bytes) -> str:
            self.put_calls += 1
            if self.fail_writes:
                raise OSError("object store unavailable")
            return super().put_bytes(key, data)

    settings = Settings(
        app_mode="production",
        database_url=get_settings().database_url,
        checkpoint_database_url=get_settings().checkpoint_database_url,
        real_model_enabled=False,
        admin_token=PRODUCTION_ADMIN_TOKEN,
        object_store_type="s3",
        s3_endpoint="http://minio:9000",
        s3_region="us-east-1",
        s3_bucket="bidscope-test",
        s3_access_key="test-access",
        s3_secret_key="test-secret",
        allowed_origins=["https://bidscope.test"],
        trusted_hosts=["bidscope.test"],
        external_scheme="https",
        object_store_root=str(tmp_path / "objects"),
    )
    headers = {"X-Admin-Token": PRODUCTION_ADMIN_TOKEN}
    store = CountingObjectStore(tmp_path / "objects")
    with TestClient(create_app(settings=settings)) as client:
        service = client.app.state.run_service
        service.graph = object()
        service.object_store = store

        async def _persist_and_export() -> tuple[str, str, str]:
            run_id, created = await service.create_run("DOCX object recovery test")
            assert created is True
            report = Report(
                run_id=run_id,
                generated_at=datetime(2026, 7, 18, 12, 0, tzinfo=UTC),
                query_conditions={"region": "四川"},
            )
            persistence = ReportPersistence(service.session_factory, store)
            persisted = await persistence.persist_online_report(report, {})
            export = await persistence.export_docx(persisted)
            return run_id, persisted.id, export.object_key

        assert client.portal is not None
        run_id, report_id, object_key = client.portal.call(_persist_and_export)
        assert store.exists(object_key)
        assert store.put_calls == 1
        store._resolve(object_key).unlink()
        assert not store.exists(object_key)
        store.fail_writes = True

        failed_retry = client.post(f"/api/reports/{run_id}/docx/retry", headers=headers)
        assert failed_retry.status_code == 503, failed_retry.text
        assert not store.exists(object_key)

        store.fail_writes = False
        retry = client.post(f"/api/reports/{run_id}/docx/retry", headers=headers)
        assert retry.status_code == 200, retry.text
        assert retry.json() == {"report_id": report_id, "docx_object_key": object_key}
        assert store.exists(object_key)
        assert store.put_calls == 3

        download = client.get(f"/api/reports/{run_id}/docx", headers=headers)
        assert download.status_code == 200, download.text
        assert download.content == store.get_bytes(object_key)
        assert isinstance(docx.Document(io.BytesIO(download.content)), docx.document.Document)

        idempotent = client.post(f"/api/reports/{run_id}/docx/retry", headers=headers)
        assert idempotent.status_code == 200, idempotent.text
        assert idempotent.json() == retry.json()
        assert store.put_calls == 3

        async def _stored_key() -> str | None:
            async with service.session_factory() as session:
                row = await session.get(ReportModel, report_id)
                assert row is not None
                return row.docx_object_key

        assert client.portal.call(_stored_key) == object_key


def test_retry_succeeds_when_retryable(demo_client: TestClient) -> None:
    """A retryable run can be retried, which re-executes the graph."""
    client = demo_client

    # Insert a retryable run directly (mark_stale_runs_retryable's output shape).
    from bidscope.db import create_engine_and_session

    _, session_factory = create_engine_and_session()

    async def _insert() -> str:
        async with session_factory() as session:
            run = QueryRun(
                run_key="retry-seed",
                status="retryable",
                user_request=NON_SCHEDULED_QUERY,
                checkpoint_thread_id="retry-seed",
            )
            session.add(run)
            await session.commit()
            await session.refresh(run)
            return run.id

    run_id = asyncio.run(_insert())

    response = client.post(f"/api/runs/{run_id}/retry")
    assert response.status_code == 200, response.text

    final = _poll_status(client, run_id, "completed")
    assert final["status"] == "completed"
