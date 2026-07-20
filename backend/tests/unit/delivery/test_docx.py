"""Unit tests for the DOCX report renderer.

These tests verify that :func:`bidscope.delivery.docx.render_report` converts a
typed :class:`~bidscope.domain.reports.Report` into a valid DOCX byte stream
that contains every element a reader relies on: the query conditions, each item
title, an unknown-field marker, the source URL, evidence labels, and a
completeness warning.

The renderer is a pure function: it touches no database, object store, or
network. That keeps the unit tests fast and fully offline.
"""

from __future__ import annotations

import io
from datetime import UTC, datetime

import docx
from bidscope.domain.reports import (
    Report,
    ReportCitation,
    ReportClaim,
    ReportItem,
)

SOURCE_URL = "https://example.invalid/demo/001"


def _sample_report(with_warning: bool = True) -> Report:
    """Build a representative synthetic-demo report for rendering tests."""
    return Report(
        run_id="demo-run-001",
        generated_at=datetime(2026, 7, 18, 12, 0, 0, tzinfo=UTC),
        query_conditions={"region": "四川", "budget": "≥500万", "topic": "智算中心"},
        freshness_window="2026-07-11 to 2026-07-18",
        source_availability=["ccgp"],
        completeness_warning=(
            "部分数据源暂不可用，结果可能不完整。" if with_warning else None
        ),
        items=[
            ReportItem(
                notice_id="demo-001",
                title="四川智算中心服务器采购招标公告",
                known_fields={
                    "source_url": SOURCE_URL,
                    "purchaser": "四川省某单位",
                    "budget": "CNY 8,000,000",
                },
                unknown_fields=["deadline"],
                relevance_reason="Matches Sichuan + computing center + budget filter.",
                risk_note=None,
                citations=[
                    ReportCitation(evidence_id="ev-001", label="预算金额证据"),
                    ReportCitation(evidence_id="ev-002", label="采购人信息"),
                ],
                claims=[
                    ReportClaim(
                        text="预算800万元，超过500万阈值。",
                        citation_ids=["ev-001"],
                    ),
                ],
            ),
            ReportItem(
                notice_id="demo-002",
                title="重庆数据中心建设项目",
                known_fields={"source_url": "https://example.invalid/demo/002"},
                unknown_fields=[],
                relevance_reason=None,
                risk_note="预算字段缺失。",
                citations=[ReportCitation(evidence_id="ev-003", label=None)],
                claims=[],
            ),
        ],
    )


def _render_to_docx(report: Report) -> docx.Document:
    """Render the report and reopen the bytes with python-docx."""
    from bidscope.delivery.docx import render_report

    data = render_report(report)
    return docx.Document(io.BytesIO(data))


def _all_text(document: docx.Document) -> str:
    """Collect paragraph and table-cell text into one searchable blob."""
    parts: list[str] = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_render_produces_parsable_docx() -> None:
    """The byte stream must open as a valid DOCX document."""
    document = _render_to_docx(_sample_report())
    assert isinstance(document, docx.document.Document)


def test_render_includes_query_conditions() -> None:
    """Every query condition value must appear in the rendered document."""
    text = _all_text(_render_to_docx(_sample_report()))
    assert "四川" in text
    assert "≥500万" in text
    assert "智算中心" in text


def test_render_includes_each_item_title() -> None:
    """Every item title must appear in the rendered document."""
    text = _all_text(_render_to_docx(_sample_report()))
    assert "四川智算中心服务器采购招标公告" in text
    assert "重庆数据中心建设项目" in text


def test_render_marks_unknown_fields() -> None:
    """The unknown-field marker for partially-known items must be present."""
    text = _all_text(_render_to_docx(_sample_report()))
    assert "deadline" in text


def test_render_includes_source_url() -> None:
    """Source URLs from item known-fields must appear in the document."""
    text = _all_text(_render_to_docx(_sample_report()))
    assert SOURCE_URL in text
    assert "https://example.invalid/demo/002" in text


def test_render_includes_evidence_labels() -> None:
    """Citation evidence labels must be rendered."""
    text = _all_text(_render_to_docx(_sample_report()))
    assert "预算金额证据" in text
    assert "采购人信息" in text


def test_render_includes_completeness_warning() -> None:
    """A non-null completeness warning must be rendered."""
    text = _all_text(_render_to_docx(_sample_report(with_warning=True)))
    assert "部分数据源暂不可用" in text


def test_render_without_warning_omits_warning_text() -> None:
    """When there is no completeness warning, the warning text is omitted."""
    text = _all_text(_render_to_docx(_sample_report(with_warning=False)))
    assert "部分数据源暂不可用" not in text


def test_render_is_deterministic() -> None:
    """Rendering the same report twice yields identical bytes."""
    from bidscope.delivery.docx import render_report

    report = _sample_report()
    assert render_report(report) == render_report(report)
